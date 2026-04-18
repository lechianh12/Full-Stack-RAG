import asyncio
import csv
import io
import os
import subprocess
import tempfile
from pathlib import Path

import docx
from fastapi import UploadFile

_MINERU_TIMEOUT = 300
_MD_CACHE_DIR = Path(os.getenv("MD_CACHE_DIR", "/app/md_cache"))


def _cache_path(filename: str) -> Path:
    stem = Path(filename).stem
    return _MD_CACHE_DIR / f"{stem}.md"


def _run_mineru(pdf_bytes: bytes, filename: str) -> str:
    cache_file = _cache_path(filename)

    if cache_file.exists():
        return cache_file.read_text(encoding="utf-8")

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        pdf_path = tmp / filename
        pdf_path.write_bytes(pdf_bytes)
        output_dir = tmp / "output"
        output_dir.mkdir()

        result = subprocess.run(
            ["mineru", "-p", str(pdf_path), "-o", str(output_dir), "-b", "pipeline", "-m", "auto"],
            capture_output=True,
            text=True,
            timeout=_MINERU_TIMEOUT,
        )
        if result.returncode != 0:
            raise RuntimeError(f"mineru failed:\n{result.stderr}")

        md_files = list(output_dir.rglob("*.md"))
        if not md_files:
            raise FileNotFoundError("mineru không tạo ra file .md")

        md_text = md_files[0].read_text(encoding="utf-8")

        _MD_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        cache_file.write_text(md_text, encoding="utf-8")

        return md_text


class DocumentExtractor:
    @staticmethod
    async def extract_text_from_pdf(file: UploadFile) -> str:
        pdf_bytes = await file.read()
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _run_mineru, pdf_bytes, file.filename or "doc.pdf")

    @staticmethod
    async def extract_text_from_csv(file: UploadFile) -> str:
        contents = await file.read()
        reader = csv.reader(io.StringIO(contents.decode("utf-8")))
        return "\n".join(", ".join(row) for row in reader)

    @staticmethod
    async def extract_text_from_docx(file: UploadFile) -> str:
        contents = await file.read()
        doc = docx.Document(io.BytesIO(contents))
        return "\n".join(para.text for para in doc.paragraphs)

    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        ext = file.filename.split(".")[-1].lower()
        if ext == "pdf":
            return await DocumentExtractor.extract_text_from_pdf(file)
        elif ext == "csv":
            return await DocumentExtractor.extract_text_from_csv(file)
        elif ext == "docx":
            return await DocumentExtractor.extract_text_from_docx(file)
        raise ValueError(f"Unsupported file extension: {ext}")


class DocumentExtractor_2:
    def extract_text_from_pdf(self, path: str | Path) -> str:
        return _run_mineru(Path(path).read_bytes(), Path(path).name)

    def extract_text_from_csv(self, path: str | Path) -> str:
        with open(path, encoding="utf-8") as f:
            return "\n".join(", ".join(row) for row in csv.reader(f))

    def extract_text_from_docx(self, path: str | Path) -> str:
        return "\n".join(para.text for para in docx.Document(path).paragraphs)

    def extract_text(self, path: str | Path) -> str:
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            return self.extract_text_from_pdf(path)
        elif ext == ".csv":
            return self.extract_text_from_csv(path)
        elif ext == ".docx":
            return self.extract_text_from_docx(path)
        raise ValueError(f"Unsupported file extension: {ext}")
